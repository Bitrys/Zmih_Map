"""
This application is a training work on WEB-7 " WEB. Знакомство с API" and pretends that this is a fucking cool map (no).
"""

from docx import Document
from docx.shared import Inches
from PyQt5.QtWidgets import *
from PyQt5.QtGui import QPixmap
from PyQt5 import QtCore
from PyQt5 import uic
from PIL import Image
from io import BytesIO

import os
import sys
import requests
import datetime
import configparser


class Settings(QWidget):
    def __init__(self):
        """
        Initialized a program.
        :return: no returns
        """
        super(Settings, self).__init__()
        uic.loadUi('setting.ui', self)

        self.saved.clicked.connect(self.save_settings)

        config = configparser.ConfigParser()
        config.read('config.ini')

        self.scale.setText(config['Geogrpahic']['spin'])
        self.d_scale.setText(config['Geogrpahic']['delta_spin'])
        self.ID_map.setText(config['Geogrpahic']['id_type_map'])
        self.mail.setText(config['Geogrpahic']['mail_point'])
        self.latit.setText(config['Geogrpahic']['latitude'])
        self.longit.setText(config['Geogrpahic']['longitude'])

    def save_settings(self):
        config = configparser.ConfigParser()
        config.read('config.ini')

        config.set('Geogrpahic', 'spin', self.scale.text())
        config.set('Geogrpahic', 'delta_spin', self.d_scale.text())
        config.set('Geogrpahic', 'id_type_map', self.ID_map.text())
        config.set('Geogrpahic', 'mail_point', self.mail.text())
        config.set('Geogrpahic', 'latitude', self.latit.text())
        config.set('Geogrpahic', 'longitude', self.longit.text())

        with open('config.ini', "w") as config_file:
            config.write(config_file)

        self.close()


class Main(QMainWindow):
    def __init__(self):
        """
        Initialized a program.
        :return: no returns
        """
        super(Main, self).__init__()
        uic.loadUi('map.ui', self)

        # Костыль
        self.type_of_map.setEditable(True)
        self.type_of_map.lineEdit().setAlignment(QtCore.Qt.AlignCenter)
        self.type_of_map.setEditable(False)

        # setup map and connect buttons
        self.reset_map()
        self.findit.clicked.connect(self.set_map)
        self.resetit.clicked.connect(self.reset_map)
        self.printresult.clicked.connect(self.print_map)
        self.exit_button.clicked.connect(self.quit_programm)
        self.mail_address.clicked.connect(self.set_map)
        self.settings_button.clicked.connect(self.setup_menu)
        self.deleted_menu.triggered.connect(self.deleted_temp)
        self.type_of_map.currentIndexChanged.connect(self.set_map)

    def quit_programm(self):
        sys.exit(1)

    def set_map(self):
        """
        Setup map.
        :return: no returns
        """
        self.clear_output()
        self.api_req()

        map_photo = QPixmap(self.map_file)
        map_photo = map_photo.scaled(650, 450)
        self.map_line.setPixmap(map_photo)

        self.get_address()

    def reset_map(self):
        """
        Reset all config and search result.
        :return: no returns
        """
        config = configparser.ConfigParser()
        config.read('config.ini')

        self.spin.setValue(float(config['Geogrpahic']['spin']))
        self.to_step.setValue(float(config['Geogrpahic']['delta_spin']))
        self.type_of_map.setCurrentIndex(int(config['Geogrpahic']['id_type_map']))
        self.mail_address.setChecked(bool(int(config['Geogrpahic']['mail_point'])))
        self.latit_inp.setText(config['Geogrpahic']['latitude'])
        self.longit_inp.setText(config['Geogrpahic']['longitude'])
        self.point_to_find.setText('')
        self.clear_output()
        self.set_map()

    def clear_output(self):
        """
        Clear output date.
        :return: no returns
        """
        self.post_index_full.setText('')
        self.full_address.setText('')

    def api_req(self):
        """
        Generates an API request from the available input and generates a map to the output.
        :return: no returns
        """
        if self.point_to_find.text() == '':
            latitude = self.latit_inp.text()
            longitude = self.longit_inp.text()
            spn = self.spin.value()
        else:
            point = self.point_to_find.text()
            API_request = f'http://geocode-maps.yandex.ru/1.x/?apikey=40d1649f-0493-4b70-98ba-98533de7710b&geocode=' \
                          f'{point}&format=json'
            response = requests.get(API_request)

            if not response:
                print('Ошибка выполнения запроса:')
                print(API_request)
                print('Http статус:', response.status_code, '(', response.reason, ')')
                sys.exit(1)
            else:
                try:
                    json_response = response.json()
                    toponym = json_response['response']['GeoObjectCollection']['featureMember'][0]['GeoObject']
                    toponym_coodrinates = toponym['Point']['pos']
                    latitude, longitude = toponym_coodrinates.split(' ')
                    self.latit_inp.setText(latitude)
                    self.longit_inp.setText(longitude)
                    spn = self.spin.value()
                except:
                    self.full_address.setText('Адрес не существует')

        types = {
            'Режим «Карта»': 'map',
            'Режим «Спутник»': 'sat',
            'Режим «Гибрид»': 'sat,skl'
        }

        API_request = f'https://static-maps.yandex.ru/1.x/?ll={latitude},{longitude}&spn={spn},{spn}&l=' \
                      f'{types[self.type_of_map.currentText()]}&size=650,450&pt={latitude},{longitude},flag'
        response = requests.get(API_request)

        if not response:
            print('Ошибка выполнения запроса:')
            print(API_request)
            print('Http статус:', response.status_code, '(', response.reason, ')')
            sys.exit(1)
        else:
            self.map_file = 'temp/map.png'
            image = Image.open(BytesIO(response.content))
            image.save('temp/map.png')

    def get_address(self):
        """
        Get address
        :return: no returns
        """

        if self.point_to_find.text() == '':
            latitude = self.latit_inp.text()
            longitude = self.longit_inp.text()
            API_request = f'http://geocode-maps.yandex.ru/1.x/?apikey=40d1649f-0493-4b70-98ba-98533de7710b&geocode=' \
                          f'{latitude},{longitude}&format=json'
            response = requests.get(API_request)
        else:
            point = self.point_to_find.text()
            API_request = f'http://geocode-maps.yandex.ru/1.x/?apikey=40d1649f-0493-4b70-98ba-98533de7710b&geocode=' \
                          f'{point}&format=json'
            response = requests.get(API_request)

        if not response:
            print('Ошибка выполнения запроса:')
            print(API_request)
            print('Http статус:', response.status_code, '(', response.reason, ')')
            sys.exit(1)
        else:
            json_response = response.json()

            try:
                toponym = json_response['response']['GeoObjectCollection']['featureMember'][0]['GeoObject']
                toponym_address = toponym['metaDataProperty']['GeocoderMetaData']['text']

                if self.mail_address.isChecked():
                    try:
                        post_code = toponym["metaDataProperty"]["GeocoderMetaData"]["Address"]["postal_code"]

                        self.full_address.setText(f'Полный адрес места: {toponym_address}')
                        self.post_index_full.setText(f'Почтовый индекс: {post_code}')
                    except:
                        self.full_address.setText(f'Полный адрес места: {toponym_address}')
                        self.post_index_full.setText('Почтовый индекс недоступен в данном регионе, или не существует')
                else:
                    self.full_address.setText(f'Полный адрес места: {toponym_address}')
            except:
                self.full_address.setText('Error!')

    def print_map(self):
        """
        Printed and saved map at docx
        :return: no returns
        """
        self.api_req()
        map_photo = QPixmap(self.map_file)

        document = Document()
        now = datetime.datetime.now()
        document.add_heading(f'Результат поиска от {now.strftime("%d-%m-%Y %H:%M")}', 0)
        document.add_heading('Карта запроса:', level=1)
        document.add_picture('temp/map.png', width=Inches(6))
        document.add_heading('Текущий адрес:', level=1)

        if self.point_to_find.text() == '':
            latitude = self.latit_inp.text()
            longitude = self.longit_inp.text()
            API_request = f'http://geocode-maps.yandex.ru/1.x/?apikey=40d1649f-0493-4b70-98ba-98533de7710b&geocode=' \
                          f'{latitude},{longitude}&format=json'
            response = requests.get(API_request)
        else:
            point = self.point_to_find.text()
            API_request = f'http://geocode-maps.yandex.ru/1.x/?apikey=40d1649f-0493-4b70-98ba-98533de7710b&geocode=' \
                          f'{point}&format=json'
            response = requests.get(API_request)

        json_response = response.json()
        toponym = json_response['response']['GeoObjectCollection']['featureMember'][0]['GeoObject']
        document.add_paragraph(toponym['metaDataProperty']['GeocoderMetaData']['text'])
        document.add_heading('Координаты места:', 1)
        document.add_paragraph(toponym['Point']['pos'])
        try:
            post_code = toponym["metaDataProperty"]["GeocoderMetaData"]["Address"]["postal_code"]
            document.add_heading('Почтовый индекс:', 1)
            document.add_paragraph(post_code)
        except:
            pass

        document.save('result.docx')

    @staticmethod
    def deleted_temp():
        """
        Deleted temporary files.
        :return:
        """
        os.remove('temp/map.png')

    def setup_menu(self):
        """
        Setting this program.
        :return:
        """
        self.Settings = Settings()
        self.Settings.show()

    def keyPressEvent(self, event):
        if event.key() == QtCore.Qt.Key_PageUp:
            spn = self.spin.setValue(self.spin.value() + self.to_step.value())
            self.set_map()
        elif event.key() == QtCore.Qt.Key_PageDown:
            spn = self.spin.setValue(self.spin.value() - self.to_step.value())
            self.set_map()
        elif event.key() == QtCore.Qt.Key_Enter:
            self.set_map()
        elif event.key() == QtCore.Qt.Key_Escape:
            sys.exit(1)


def except_hook(cls, exception, traceback):
    sys.__excepthook__(cls, exception, traceback)


if __name__ == "__main__":
    app = QApplication(sys.argv)
    ex = Main()
    ex.show()
    sys.excepthook = except_hook
    sys.exit(app.exec())
